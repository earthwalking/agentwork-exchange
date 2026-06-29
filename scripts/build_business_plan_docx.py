from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "business-plan-v2.md"
OUTPUT = ROOT / "docs" / "AgentWork_Exchange_Business_Plan_V2.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(32, 37, 44)
MUTED = RGBColor(100, 113, 125)
LIGHT_FILL = "F4F6F9"
HEADER_FILL = "E8EEF5"
BORDER = "DADCE0"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_style_font(style, size, color=INK, bold=False):
    style.font.name = "Microsoft YaHei"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_paragraph_spacing(paragraph, before=0, after=8, line=1.333):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_paragraph(doc, text, style=None, bold_prefix=False):
    paragraph = doc.add_paragraph(style=style)
    set_paragraph_spacing(paragraph)
    if bold_prefix and "：" in text:
      prefix, rest = text.split("：", 1)
      run = paragraph.add_run(prefix + "：")
      set_run_font(run, bold=True)
      run = paragraph.add_run(rest)
      set_run_font(run)
    else:
      run = paragraph.add_run(text)
      set_run_font(run)
    return paragraph


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=48, after=12)
    r = p.add_run("AgentWork Exchange")
    set_run_font(r, size=28, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=6)
    r = p.add_run("商业计划书 V2")
    set_run_font(r, size=20, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=24)
    r = p.add_run("基于个人 Agent 的赏金任务采集、撮合、交付与结算平台")
    set_run_font(r, size=12, color=MUTED)

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    set_table_widths(table, [1.7, 4.5])
    rows = [
        ("版本日期", "2026-06-29"),
        ("核心场景", "GitHub 与开源社区赏金任务"),
        ("首批 Agent", "Hermes Bounty Scout / OpenAI Codex Bounty Implementer"),
        ("交付形态", "开源 Web MVP + CLI + 商业计划书"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        set_cell_shading(row.cells[0], HEADER_FILL)
        for idx, text in enumerate([label, value]):
            row.cells[idx].text = ""
            p = row.cells[idx].paragraphs[0]
            set_paragraph_spacing(p, after=0, line=1.15)
            r = p.add_run(text)
            set_run_font(r, size=10.5, bold=idx == 0, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=18, after=8)
    r = p.add_run("Prepared for open-source release and early commercial validation")
    set_run_font(r, size=10.5, color=MUTED)
    doc.add_page_break()


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip().strip("|")
        cells = [cell.strip() for cell in line.split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_table(doc, rows):
    if not rows:
        return
    col_count = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    widths = {
        2: [2.0, 4.3],
        3: [1.55, 2.35, 2.35],
        4: [1.4, 1.2, 2.7, 1.0],
    }.get(col_count, [6.3 / col_count] * col_count)
    set_table_widths(table, widths)
    for row_idx, row in enumerate(rows):
        for col_idx, text in enumerate(row):
            cell = table.rows[row_idx].cells[col_idx]
            if row_idx == 0:
                set_cell_shading(cell, HEADER_FILL)
            cell.text = ""
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, after=0, line=1.15)
            r = p.add_run(text)
            set_run_font(r, size=9.2, bold=row_idx == 0, color=INK)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, after=4)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    set_style_font(styles["Normal"], 11, INK)
    styles["Normal"].paragraph_format.space_after = Pt(8)
    styles["Normal"].paragraph_format.line_spacing = 1.333
    set_style_font(styles["Heading 1"], 16, BLUE, True)
    styles["Heading 1"].paragraph_format.space_before = Pt(18)
    styles["Heading 1"].paragraph_format.space_after = Pt(10)
    set_style_font(styles["Heading 2"], 13, BLUE, True)
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    set_style_font(styles["Heading 3"], 12, DARK_BLUE, True)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)

    add_cover(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=2)
            i += 1
            continue
        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            set_paragraph_spacing(p, after=4, line=1.208)
            r = p.add_run(line[2:])
            set_run_font(r, size=10.5)
            i += 1
            continue
        if re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            set_paragraph_spacing(p, after=4, line=1.208)
            r = p.add_run(re.sub(r"^\d+\. ", "", line))
            set_run_font(r, size=10.5)
            i += 1
            continue
        if line.startswith("版本日期") or line.startswith("项目定位"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing(p, after=2, line=1.15)
            r = p.add_run(line)
            set_run_font(r, size=9.5, color=MUTED)
            i += 1
            continue
        add_paragraph(doc, line, bold_prefix=True)
        i += 1

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("AgentWork Exchange Business Plan V2")
    set_run_font(run, size=8.5, color=MUTED)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
